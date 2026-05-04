import os
import subprocess
import sys
import re
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QPushButton,
    QFileDialog, QMessageBox, QLineEdit, QHBoxLayout, QProgressBar,
    QTabWidget, QGroupBox
)
from PyQt5.QtCore import Qt, QUrl                    # NUEVO: QUrl para enlace
from PyQt5.QtGui import QFont, QIcon, QDesktopServices  # NUEVO: QDesktopServices

# Conversiones base
from pdf2docx import Converter
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx2pdf import convert as docx2pdf_convert
from PIL import Image

OCR_LANG = "spa"

# ------------------------------------------------------------
# Funciones de conversión (mantenidas y mejoradas)
# ------------------------------------------------------------

def pdf_has_text(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            if page.extract_text():
                return True
        return False
    except Exception:
        return False

def convert_with_pdf2docx(pdf_path, word_path):
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

def convert_with_ocr(pdf_path, word_path):
    pages = convert_from_path(pdf_path, dpi=150)
    doc = Document()
    for page_num, page in enumerate(pages, start=1):
        text = pytesseract.image_to_string(page, lang=OCR_LANG)
        doc.add_paragraph(f"--- Pagina {page_num} ---")
        doc.add_paragraph(text if text.strip() else "[Sin texto detectado]")
    doc.save(word_path)

def convert_word_to_pdf(word_path, pdf_path):
    """
    Convierte DOCX a PDF usando primero Microsoft Word, y si falla, prueba con LibreOffice.
    Lanza una excepción si no encuentra ninguna opción.
    """
    # --- Intento 1: Usar Microsoft Word ---
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(word_path, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            return  # Conversión exitosa con Word
        else:
            raise Exception("El archivo PDF generado por Word está vacío o no se creó.")
    except Exception as e_word:
        # Si falla, lo registramos y pasamos al siguiente método
        print(f"Advertencia: Falló la conversión con Microsoft Word: {e_word}")
    
    # --- Intento 2: Usar LibreOffice (si está instalado) ---
    try:
        # Buscar la ruta de ejecución de LibreOffice
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        soffice_cmd = None
        for path in libreoffice_paths:
            if os.path.exists(path):
                soffice_cmd = path
                break
        
        # Si no se encontró, buscamos en el PATH del sistema
        if soffice_cmd is None:
            result = subprocess.run(['where', 'soffice'], capture_output=True, text=True)
            if result.returncode == 0:
                soffice_cmd = result.stdout.strip().split('\n')[0]
        
        if soffice_cmd is None:
            raise Exception("No se encontró LibreOffice en el sistema. Instala LibreOffice desde https://libreoffice.org")
        
        # Ejecutar LibreOffice en modo headless para la conversión
        cmd = [soffice_cmd, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), word_path]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        # LibreOffice guarda el PDF en el directorio de salida con el mismo nombre base
        generated_pdf = os.path.join(os.path.dirname(pdf_path), os.path.basename(word_path).replace('.docx', '.pdf'))
        if os.path.exists(generated_pdf):
            # Si el archivo generado no está en la ruta exacta que queríamos, lo movemos
            if generated_pdf != pdf_path:
                os.rename(generated_pdf, pdf_path)
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return  # Conversión exitosa con LibreOffice
            else:
                raise Exception("El archivo PDF generado por LibreOffice está vacío.")
        else:
            raise Exception(f"No se pudo generar el PDF. Salida de LibreOffice: {result.stderr}")
            
    except subprocess.TimeoutExpired:
        raise Exception("La conversión con LibreOffice excedió el tiempo límite de 60 segundos.")
    except Exception as e_libre:
        # Si ambos métodos fallan, mostramos un mensaje claro
        raise Exception(
            f"No se pudo convertir el archivo Word a PDF.\n"
            f"Razón: {str(e_libre)}\n\n"
            f"Soluciones posibles:\n"
            f"1. Instala Microsoft Word (incluido en Microsoft Office).\n"
            f"2. Instala LibreOffice (gratuito) desde https://libreoffice.org\n"
            f"3. Verifica que el archivo '{os.path.basename(word_path)}' no esté dañado."
        )

def convert_word_to_txt(word_path, txt_path):
    doc = Document(word_path)
    with open(txt_path, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

def convert_txt_to_word(txt_path, word_path):
    doc = Document()
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                doc.add_paragraph(line.strip())
    doc.save(word_path)

def convert_images_to_pdf(images_folder, output_pdf):
    """
    Convierte todas las imágenes de una carpeta a un solo PDF.
    Ordena los archivos numéricamente por el número en el nombre.
    """
    valid_ext = ('.png', '.jpg', '.jpeg', '.bmp')
    images = []
    for f in os.listdir(images_folder):
        if f.lower().endswith(valid_ext):
            images.append(os.path.join(images_folder, f))
    if not images:
        raise ValueError("No se encontraron imágenes en la carpeta")

    def extract_number(filename):
        numbers = re.findall(r'\d+', filename)
        return int(numbers[0]) if numbers else 0

    images.sort(key=lambda x: extract_number(os.path.basename(x)))

    image_list = []
    for img_path in images:
        img = Image.open(img_path)
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        image_list.append(img)

    if image_list:
        image_list[0].save(output_pdf, save_all=True, append_images=image_list[1:])
    else:
        raise ValueError("No hay imágenes válidas para convertir")

# ------------------------------------------------------------
# Interfaz gráfica profesional
# ------------------------------------------------------------

class ConverterApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Convertidor Profesional de Documentos Hailie v2.1")  # MODIFICADO: v2.1
        self.setGeometry(150, 150, 750, 550)
        self.setMinimumSize(700, 500)

        # Estilos globales (profesional, sin emojis)
        self.setStyleSheet("""
            QWidget {
                background-color: #f0f2f5;
                font-family: "Segoe UI", Arial, sans-serif;
                font-size: 12px;
            }
            QGroupBox {
                font-weight: bold;
                border: 1px solid #c0c0c0;
                border-radius: 5px;
                margin-top: 12px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
            QPushButton {
                background-color: #2c3e50;
                color: white;
                border: none;
                padding: 8px 12px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1e2b36;
            }
            QPushButton:disabled {
                background-color: #95a5a6;
            }
            QLineEdit {
                padding: 5px;
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                background-color: white;
            }
            QLineEdit:focus {
                border: 1px solid #2c3e50;
            }
            QProgressBar {
                border: 1px solid #bdc3c7;
                border-radius: 3px;
                text-align: center;
                background-color: white;
            }
            QProgressBar::chunk {
                background-color: #27ae60;
                width: 10px;
            }
            QTabWidget::pane {
                border: 1px solid #c0c0c0;
                background: white;
            }
            QTabBar::tab {
                background: #e0e0e0;
                padding: 6px 12px;
                margin: 2px;
            }
            QTabBar::tab:selected {
                background: #2c3e50;
                color: white;
            }
            QLabel#status_label {
                padding: 5px;
                color: #2c3e50;
                background-color: #ecf0f1;
                border-radius: 3px;
            }
        """)

        layout = QVBoxLayout()

        # Título principal
        title = QLabel("Convertidor Profesional de Documentos - Hailie v2.1")  # MODIFICADO: v2.1
        title.setAlignment(Qt.AlignCenter)
        title.setFont(QFont("Segoe UI", 18, QFont.Bold))
        title.setStyleSheet("color: #2c3e50; margin: 10px;")
        layout.addWidget(title)

        # Grupo de carpetas
        folder_group = QGroupBox("Configuración de carpetas")
        folder_layout = QVBoxLayout()

        hbox_in = QHBoxLayout()
        self.input_edit = QLineEdit()
        self.input_edit.setPlaceholderText("Seleccione la carpeta de entrada")
        btn_input = QPushButton("Examinar...")
        btn_input.clicked.connect(self.select_input)
        hbox_in.addWidget(QLabel("Entrada:"))
        hbox_in.addWidget(self.input_edit)
        hbox_in.addWidget(btn_input)

        hbox_out = QHBoxLayout()
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("Seleccione la carpeta de salida")
        btn_output = QPushButton("Examinar...")
        btn_output.clicked.connect(self.select_output)
        hbox_out.addWidget(QLabel("Salida:"))
        hbox_out.addWidget(self.output_edit)
        hbox_out.addWidget(btn_output)

        folder_layout.addLayout(hbox_in)
        folder_layout.addLayout(hbox_out)
        folder_group.setLayout(folder_layout)
        layout.addWidget(folder_group)

        # Barra de progreso y estado
        self.progress = QProgressBar()
        layout.addWidget(self.progress)
        self.status_label = QLabel("Listo")
        self.status_label.setObjectName("status_label")
        self.status_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.status_label)

        # Pestañas
        self.tabs = QTabWidget()
        self.tabs.addTab(self.create_pdf_word_tab(), "PDF <-> Word")
        self.tabs.addTab(self.create_word_txt_tab(), "Word <-> TXT")
        self.tabs.addTab(self.create_img_to_pdf_tab(), "Imagen -> PDF")
        layout.addWidget(self.tabs)

        # Botones inferiores (Ayuda, Acerca de, Donación)
        hbox_bottom = QHBoxLayout()
        btn_help = QPushButton("Ayuda")
        btn_help.setStyleSheet("background-color: #f39c12;")
        btn_help.clicked.connect(self.show_help)
        
        btn_about = QPushButton("Acerca de")
        btn_about.setStyleSheet("background-color: #3498db;")
        btn_about.clicked.connect(self.show_about)
        
        # NUEVO: Botón de donación
        btn_donate = QPushButton("Donar (PayPal)")
        btn_donate.setStyleSheet("background-color: #e67e22; color: white;")
        btn_donate.clicked.connect(self.open_donation_link)
        
        hbox_bottom.addStretch()
        hbox_bottom.addWidget(btn_help)
        hbox_bottom.addWidget(btn_about)
        hbox_bottom.addWidget(btn_donate)   # NUEVO
        layout.addLayout(hbox_bottom)

        self.setLayout(layout)

    # --------------------------------------------------------
    # Métodos auxiliares
    # --------------------------------------------------------
    def select_input(self):
        folder = QFileDialog.getExistingDirectory(self, "Carpeta de entrada")
        if folder:
            self.input_edit.setText(folder)

    def select_output(self):
        folder = QFileDialog.getExistingDirectory(self, "Carpeta de salida")
        if folder:
            self.output_edit.setText(folder)

    def validate_folders(self):
        inp = self.input_edit.text().strip()
        outp = self.output_edit.text().strip()
        if not inp or not os.path.isdir(inp):
            QMessageBox.warning(self, "Error", "Carpeta de entrada no válida.")
            return False, None, None
        if not outp or not os.path.isdir(outp):
            QMessageBox.warning(self, "Error", "Carpeta de salida no válida.")
            return False, None, None
        return True, inp, outp

    def update_progress(self, value, text=None):
        self.progress.setValue(value)
        if text is not None:
            self.status_label.setText(text)
        QApplication.processEvents()

    # NUEVO: Abrir enlace de donación
    def open_donation_link(self):
        url = QUrl("https://paypal.me/danievilmathers?locale.x=es_XC&country.x=MX")
        QDesktopServices.openUrl(url)

    # --------------------------------------------------------
    # Creación de pestañas
    # --------------------------------------------------------
    def create_pdf_word_tab(self):
        widget = QWidget()
        layout = QVBoxLayout()
        btn_pdf_to_word = QPushButton("Convertir PDF a Word (múltiples archivos)")
        btn_pdf_to_word.clicked.connect(self.run_pdf_to_word)
        btn_word_to_pdf = QPushButton("Convertir Word a PDF (múltiples archivos)")
        btn_word_to_pdf.clicked.connect(self.run_word_to_pdf)
        layout.addWidget(btn_pdf_to_word)
        layout.addWidget(btn_word_to_pdf)
        layout.addStretch()
        widget.setLayout(layout)
        return widget

    def create_word_txt_tab(self):
        widget = QWidget()
        layout = QVBoxLayout()
        btn_word_to_txt = QPushButton("Convertir Word a TXT (múltiples archivos)")
        btn_word_to_txt.clicked.connect(self.run_word_to_txt)
        btn_txt_to_word = QPushButton("Convertir TXT a Word (múltiples archivos)")
        btn_txt_to_word.clicked.connect(self.run_txt_to_word)
        layout.addWidget(btn_word_to_txt)
        layout.addWidget(btn_txt_to_word)
        layout.addStretch()
        widget.setLayout(layout)
        return widget

    def create_img_to_pdf_tab(self):
        widget = QWidget()
        layout = QVBoxLayout()
        btn_img_to_pdf = QPushButton("Convertir imágenes de la carpeta a PDF")
        btn_img_to_pdf.clicked.connect(self.run_img_to_pdf)
        layout.addWidget(btn_img_to_pdf)
        layout.addStretch()
        widget.setLayout(layout)
        return widget

    # --------------------------------------------------------
    # Ejecutores de conversión (robustos)
    # --------------------------------------------------------
    def run_pdf_to_word(self):
        ok, inp, outp = self.validate_folders()
        if not ok:
            return
        files = [f for f in os.listdir(inp) if f.lower().endswith(".pdf")]
        if not files:
            QMessageBox.warning(self, "Aviso", "No hay archivos PDF en la carpeta de entrada.")
            return
        total = len(files)
        for i, file in enumerate(files, start=1):
            pdf_path = os.path.join(inp, file)
            word_path = os.path.join(outp, file.replace(".pdf", ".docx"))
            self.update_progress(int((i-1)/total*100), f"Procesando: {file}")
            try:
                if pdf_has_text(pdf_path):
                    convert_with_pdf2docx(pdf_path, word_path)
                else:
                    convert_with_ocr(pdf_path, word_path)
            except Exception as e:
                QMessageBox.warning(self, "Error en archivo", f"Fallo en {file}:\n{str(e)}")
        self.update_progress(100, "Conversión completada.")
        QMessageBox.information(self, "Completado", f"PDF a Word finalizado.\nSalida: {outp}")

    def run_word_to_pdf(self):
        ok, inp, outp = self.validate_folders()
        if not ok:
            return
        files = [f for f in os.listdir(inp) if f.lower().endswith(".docx")]
        if not files:
            QMessageBox.warning(self, "Aviso", "No hay archivos Word en la carpeta de entrada.")
            return
        total = len(files)
        for i, file in enumerate(files, start=1):
            word_path = os.path.join(inp, file)
            pdf_path = os.path.join(outp, file.replace(".docx", ".pdf"))
            self.update_progress(int((i-1)/total*100), f"Convirtiendo: {file}")
            try:
                convert_word_to_pdf(word_path, pdf_path)
            except Exception as e:
                QMessageBox.warning(self, "Error en archivo", f"Fallo en {file}:\n{str(e)}")
        self.update_progress(100, "Conversión completada.")
        QMessageBox.information(self, "Completado", f"Word a PDF finalizado.\nSalida: {outp}")

    def run_word_to_txt(self):
        ok, inp, outp = self.validate_folders()
        if not ok:
            return
        files = [f for f in os.listdir(inp) if f.lower().endswith(".docx")]
        if not files:
            QMessageBox.warning(self, "Aviso", "No hay archivos Word en la carpeta de entrada.")
            return
        total = len(files)
        for i, file in enumerate(files, start=1):
            word_path = os.path.join(inp, file)
            txt_path = os.path.join(outp, file.replace(".docx", ".txt"))
            self.update_progress(int((i-1)/total*100), f"Extrayendo texto: {file}")
            try:
                convert_word_to_txt(word_path, txt_path)
            except Exception as e:
                QMessageBox.warning(self, "Error en archivo", f"Fallo en {file}:\n{str(e)}")
        self.update_progress(100, "Conversión completada.")
        QMessageBox.information(self, "Completado", f"Word a TXT finalizado.\nSalida: {outp}")

    def run_txt_to_word(self):
        ok, inp, outp = self.validate_folders()
        if not ok:
            return
        files = [f for f in os.listdir(inp) if f.lower().endswith(".txt")]
        if not files:
            QMessageBox.warning(self, "Aviso", "No hay archivos TXT en la carpeta de entrada.")
            return
        total = len(files)
        for i, file in enumerate(files, start=1):
            txt_path = os.path.join(inp, file)
            word_path = os.path.join(outp, file.replace(".txt", ".docx"))
            self.update_progress(int((i-1)/total*100), f"Creando Word: {file}")
            try:
                convert_txt_to_word(txt_path, word_path)
            except Exception as e:
                QMessageBox.warning(self, "Error en archivo", f"Fallo en {file}:\n{str(e)}")
        self.update_progress(100, "Conversión completada.")
        QMessageBox.information(self, "Completado", f"TXT a Word finalizado.\nSalida: {outp}")

    def run_img_to_pdf(self):
        ok, inp, outp = self.validate_folders()
        if not ok:
            return
        valid_ext = ('.png', '.jpg', '.jpeg', '.bmp')
        has_images = any(f.lower().endswith(valid_ext) for f in os.listdir(inp))
        if not has_images:
            QMessageBox.warning(self, "Aviso", "No se encontraron imágenes (PNG, JPG, JPEG, BMP) en la carpeta de entrada.")
            return

        default_name = os.path.join(outp, "imagenes_convertidas.pdf")
        output_pdf, _ = QFileDialog.getSaveFileName(self, "Guardar PDF como", default_name, "PDF Files (*.pdf)")
        if not output_pdf:
            return

        self.update_progress(0, "Convirtiendo imágenes a PDF...")
        try:
            convert_images_to_pdf(inp, output_pdf)
            self.update_progress(100, "PDF creado exitosamente.")
            QMessageBox.information(self, "Completado", f"PDF generado:\n{output_pdf}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo crear el PDF:\n{str(e)}")
            self.update_progress(0, "Error en la conversión")

    # --------------------------------------------------------
    # Ayuda y Acerca de (actualizados)
    # --------------------------------------------------------
    def show_help(self):
        QMessageBox.information(self, "Ayuda",
            "GUIA DE USO\n\n"
            "1. Seleccione la carpeta de ENTRADA (con los archivos a convertir).\n"
            "2. Seleccione la carpeta de SALIDA (donde se guardarán los resultados).\n"
            "3. Vaya a la pestaña correspondiente.\n"
            "4. Haga clic en el botón de conversión deseado.\n\n"
            "CONVERSIONES DISPONIBLES:\n"
            "- PDF a Word: Convierte PDFs a DOCX. Si el PDF es escaneado, usa OCR automáticamente.\n"
            "- Word a PDF: Convierte DOCX a PDF (intenta primero con Microsoft Word, luego con LibreOffice).\n"
            "- Word a TXT: Extrae el texto plano de los archivos Word.\n"
            "- TXT a Word: Crea documentos Word a partir de archivos de texto.\n"
            "- Imagen a PDF: Combina todas las imágenes de una carpeta en un solo PDF (orden numérico natural).\n\n"
            "REQUISITOS ADICIONALES:\n"
            "- Para OCR: Tesseract instalado y en el PATH.\n"
            "- Para Word a PDF: Microsoft Word o LibreOffice (gratuito).")

    def show_about(self):
        QMessageBox.information(self, "Acerca de",
            "Convertidor Profesional de Documentos - Hailie Versión 2.1\n\n"
            "Desarrollado por: Daniel Donaldo Villanueva Canul\n"
            "Correo de contacto: danieviloficialcontacto@gmail.com\n\n"   # NUEVO
            "Lenguaje: Python 3 + PyQt5\n"
            "Librerías: pdf2docx, PyPDF2, pdf2image, pytesseract, python-docx, docx2pdf, Pillow\n\n"
            "Conversión masiva de documentos con interfaz profesional.\n"
            "(C) 2026 - Todos los derechos reservados.\n\n"
            "Si te es útil este software, considera hacer una donación para apoyar el desarrollo continuo.\n"
            "¡Gracias por usar Convertidor Profesional Hailie!")

# ------------------------------------------------------------
# Punto de entrada
# ------------------------------------------------------------
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ConverterApp()
    window.show()
    sys.exit(app.exec_())